"""Конвертация загруженного файла ТЗ (.docx / .md / .txt) в Markdown."""

import io
import re

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def file_to_markdown(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".docx"):
        return docx_to_markdown(data)
    if name.endswith((".md", ".markdown", ".txt")):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Неподдерживаемый формат файла: {filename}")


def docx_blocks(data: bytes) -> list[tuple[int, str]]:
    """Контентные блоки документа: (индекс среди p/tbl-детей body, markdown).

    Индекс считается по ВСЕМ абзацам/таблицам, включая пустые и не дающие
    текста (например, абзац с одной картинкой) — они индекс занимают, но в
    результат не попадают. На эти индексы опирается карта «раздел → блоки
    оригинала» (app/originals.py), по которой экспорт патчит исходный .docx.
    """
    doc = Document(io.BytesIO(data))
    out: list[tuple[int, str]] = []
    for i, block in enumerate(_iter_blocks(doc)):
        if isinstance(block, Paragraph):
            md = _paragraph_to_md(block)
        else:
            md = _table_to_md(block) or None
        if md:
            out.append((i, md))
    return out


def docx_to_markdown(data: bytes) -> str:
    return "\n\n".join(md for _, md in docx_blocks(data)).strip()


def _iter_blocks(doc):
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, doc)
        elif child.tag.endswith("}tbl"):
            yield Table(child, doc)


_HEADING_RE = re.compile(r"(?:heading|заголовок)\s*(\d)", re.IGNORECASE)
_LIST_RE = re.compile(r"list|список|маркированный|нумерованный", re.IGNORECASE)


def _paragraph_to_md(p: Paragraph) -> str | None:
    text = p.text.strip()
    if not text:
        return None
    style = p.style.name if p.style is not None else ""
    if m := _HEADING_RE.search(style):
        level = min(int(m.group(1)), 6)
        return f"{'#' * level} {text}"
    if _LIST_RE.search(style):
        return f"- {text}"
    # Стиль Word "Заголовок" применён не всегда — многие консультанты просто
    # выделяют раздел жирным. Помечаем это в Markdown, чтобы эвристика
    # нарезки на разделы (app/spec_doc.py) могла распознать границу раздела
    # и там, где реального стиля заголовка нет.
    if len(text) <= 120 and _is_all_bold(p):
        return f"**{text}**"
    return text


def _is_all_bold(p: Paragraph) -> bool:
    runs = [r for r in p.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def _table_to_md(table: Table) -> str:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([" ".join(cell.text.split()) for cell in row.cells])
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    out = ["| " + " | ".join(rows[0]) + " |", "|" + " --- |" * width]
    out += ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join(out)


# --- экспорт: Markdown документа сессии обратно в .docx для скачивания ---
# Превью на фронтенде и хранение в графе — Markdown (см. app/spec_doc.py),
# но пользователю нужен файл в исходном формате ТЗ, поэтому здесь — обратная
# конвертация. Не пытается воспроизвести весь Markdown (например, вложенные
# списки), только то, что реально встречается в ТЗ-документах: заголовки,
# абзацы с **жирным**/*курсивом*, маркированные/нумерованные списки и таблицы.

_MD_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BOLD_LINE_RE = re.compile(r"^\*\*(.+)\*\*$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^(\d+[.)])\s+(.*)$")  # маркер в группе 1 — нужен фолбэку
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def markdown_to_docx(markdown: str) -> bytes:
    doc = Document()
    append_markdown(doc, markdown)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def append_markdown(doc: Document, markdown: str) -> None:
    """Дописывает Markdown в конец документа стилями САМОГО документа.

    Используется и при сборке экспорта с нуля, и при патче чужого .docx
    (app/originals.py) — поэтому стили применяются через фолбэки: если в
    документе нет нужного стиля (Heading N, List Bullet, Table Grid),
    содержимое не теряется, а оформляется без стиля.
    """
    lines = markdown.split("\n")
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("|") and line.endswith("|"):
            block = []
            while i < n and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            _add_table(doc, block)
            continue
        if m := _MD_HEADING_RE.match(line):
            _add_heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue
        if m := _BOLD_LINE_RE.match(line):
            p = doc.add_paragraph()
            p.add_run(m.group(1).strip()).bold = True
            i += 1
            continue
        if m := _BULLET_RE.match(line):
            _add_list(doc, "List Bullet", "•", m.group(1))
            i += 1
            continue
        if m := _NUMBERED_RE.match(line):
            _add_list(doc, "List Number", m.group(1), m.group(2))
            i += 1
            continue
        para_lines = [line]
        i += 1
        while i < n and lines[i].strip():
            para_lines.append(lines[i].strip())
            i += 1
        _add_inline_runs(doc.add_paragraph(), " ".join(para_lines))


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    try:
        p.style = f"Heading {min(level, 9)}"
    except KeyError:
        p.add_run(text).bold = True
        return
    p.add_run(text)


def _add_list(doc: Document, style: str, marker: str, text: str) -> None:
    p = doc.add_paragraph()
    try:
        p.style = style
    except KeyError:
        _add_inline_runs(p, f"{marker} {text}")
        return
    _add_inline_runs(p, text)


def _add_inline_runs(paragraph, text: str) -> None:
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def _add_table(doc: Document, block_lines: list[str]) -> None:
    rows = [[c.strip() for c in ln.strip("|").split("|")] for ln in block_lines]
    if len(rows) >= 2 and all(re.fullmatch(r":?-+:?", c) for c in rows[1]):
        rows = [rows[0]] + rows[2:]
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # в документе нет такого стиля таблиц — оставляем без рамок
    for r, row in enumerate(rows):
        for c in range(width):
            cell_text = row[c] if c < len(row) else ""
            cell = table.cell(r, c)
            cell.text = cell_text
            if r == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
